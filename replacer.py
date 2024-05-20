import docx
import re
import sys
import tkinter as tk
from tkinter import ttk
from tkinter import filedialog

class ReplacerApp:
    def __init__(self, app):
        self.app = app
        self.selected_file_path = None
        self.doc = None
        self.codes = None

    def select_file(self):
        self.app.clear_error_msg()
        file = filedialog.askopenfile(mode='r')
        if file:
            self.selected_file_path = file.name
            print("File selected: ", self.selected_file_path)
            self.app.update_label("File selected: " + self.selected_file_path)

        if self.selected_file_path:
            self.open_document()
        else:
            self.app.fill_error_msg("No file selected.")

    def open_document(self):
        try:
            self.app.clear_error_msg()
            if self.selected_file_path:
                self.doc = docx.Document(self.selected_file_path)
            else:
                raise FileNotFoundError("No file selected.")
        except Exception as e:
            print(f"Error opening word document: {e}")
            self.app.fill_error_msg(f"Error opening document: {e}")
        assert self.doc is not None
        self.retrieve_codes()

    def retrieve_codes(self):
        if not self.doc:
            self.open_document()

        if not self.doc:
            self.app.fill_error_msg("No document loaded.")
            return

        self.codes = dict()
        for paragraph in self.doc.paragraphs:
            codes = re.findall(r"\{[^\}]+\}", paragraph.text)
            for code in codes:
                code = code[1:-1]
                parts = code.split("|")
                question = parts[0]
                default = parts[1] if len(parts) > 1 else ""
                self.codes[question] = default
        print(self.codes)
        self.app.create_entries()

class App(tk.Frame):
    def __init__(self, master=None):
        super().__init__(master)
        self.pack()
        self.replacer = ReplacerApp(self)
        self.create_widgets()

    def create_widgets(self):
        self.label = tk.Label(self, text="No file selected.")
        self.label.pack(pady=2)

        self.selectButton = tk.Button(self, text="Select File", command=self.replacer.select_file)
        self.selectButton.pack(pady=2)

        self.error_msg = tk.Label(self, text="")
        self.error_msg.pack(pady=5)

        self.entries_frame = tk.Frame(self)
        self.entries_frame.pack(pady=5)

    def create_entries(self):
        if not self.replacer.codes:
            return

        for (code, answer) in self.replacer.codes.items():
            label = tk.Label(self.entries_frame, text=code)
            label.pack()

            entry = tk.Entry(self.entries_frame)
            entry.pack()

            self.replacer.codes[code] = tk.StringVar()
            self.replacer.codes[code].set(answer)
            entry["textvariable"] = self.replacer.codes[code]
            entry.bind('<Key-Return>', self.print_contents)

    def print_contents(self, event):
        print("Hi. The current entry content is:", self.replacer.codes)

    def update_label(self, text):
        self.label.config(text=text)

    def fill_error_msg(self, text):
        self.error_msg.config(text=text)
    
    def clear_error_msg(self):
        self.error_msg.config(text="")

if __name__ == "__main__":
    root = tk.Tk()
    root.geometry("500x400+100+50")
    myapp = App(master=root)
    myapp.master.title("My App")
    myapp.mainloop()


"""
def prompt_user(question):
    print(question)
    answer = input()
    if answer.lower() == 'quit':
        print("Goodbye!")
        sys.exit()
    return answer

# Create a Tkinter root widget
root = tk.Tk(screenName="Format Code Replacer")
frame = ttk.Frame(root, padding=10)
frame.grid()
ttk.Label(frame, text="Hello world").grid(row=0, column=0)
ttk.Button(frame, text="Quit", command=root.destroy).grid(row=0, column=1)
root.mainloop()
#root.withdraw()  # Hide the main window

# Ask the user to select a file
file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
if not file_path:
    print("No file selected.")
    wait = input("Press enter to exit program.")
    sys.exit()

# Open the document
try:
    doc = docx.Document(file_path)
except Exception as e:
    print(f"Error opening document: {e}")
    wait = input("Press enter to exit program.")
    sys.exit()

# Replace codes in the document
answered_codes = dict() # keeps track of answered codes so repeated codes are not asked again
for paragraph in doc.paragraphs:
"""
    #codes = re.findall(r"\{[^\}]+\}", paragraph.text)
"""
    for code in codes:
        code = code[1:-1]
        parts = code.split("|")
        question = parts[0]
        default = parts[1] if len(parts) > 1 else ""
        
        answer = ""
        if question in answered_codes:
            answer = answered_codes[question]
        else:
            answer = prompt_user(question)
            if answer == "":
                answer = default
            answered_codes[question] = answer
        
        original_text = paragraph.text
        paragraph.clear()
        new_text = original_text.replace("{" + code + "}", answer)
        paragraph.add_run(new_text)
        print(answered_codes)

# Prompt user to save the modified document
try:
    output_file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if output_file_path:
        doc.save(output_file_path)
        print("The program has finished. Please check the modified document.")
    else:
        print("No save location selected. Exiting without saving.")
except Exception as e:
    print(f"Error saving document: {e}")

"""
